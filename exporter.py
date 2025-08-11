from docx import Document
from fpdf import FPDF
import html

def export_to_word(summary, articles):
    """Exports the content to a Word document."""
    document = Document()
    document.add_heading('AI News Summary', 0)

    document.add_heading('Executive Summary', level=1)
    document.add_paragraph(summary)

    document.add_heading('Articles', level=1)
    for article in articles:
        document.add_heading(article['title'], level=2)
        document.add_paragraph(f"Link: {article['link']}")
        document.add_paragraph(article['summary'])
        document.add_paragraph()

    return document

def export_to_pdf(summary, articles):
    """Exports the content to a PDF document."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt="AI News Summary", ln=True, align='C')

    pdf.ln(10)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Executive Summary", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 5, summary)

    pdf.ln(10)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Articles", ln=True)
    for article in articles:
        pdf.set_font("Arial", 'B', 12)
        pdf.multi_cell(0, 5, article['title'])
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 5, f"Link: {article['link']}")
        pdf.multi_cell(0, 5, article['summary'])
        pdf.ln(5)

    return pdf

def export_to_html(summary, articles):
    """Exports the content to an HTML string."""
    html_content = f"""<html>
<head><title>AI News Summary</title></head>
<body>
<h1>AI News Summary</h1>
<h2>Executive Summary</h2>
<p>{html.escape(summary)}</p>
<h2>Articles</h2>
"""
    for article in articles:
        html_content += f"""<h3>{html.escape(article['title'])}</h3>
<p><a href=\"{article['link']}\">{html.escape(article['link'])}</a></p>
<p>{html.escape(article['summary'])}</p>
"""
    html_content += "</body></html>"
    return html_content

